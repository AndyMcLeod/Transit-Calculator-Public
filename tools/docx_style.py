# ============================================================================
# VENDORED from the companion fuel planner: a sibling project	ools\docx_style.py
#
# Same reasoning as currents.py: a copy, not an import, because that planner is a
# separate repository and this project has to stand on its own. The trade is drift —
# this file does NOT follow its source. Diff against it before assuming a styling
# change made there applies here.
#
# It is worth copying rather than rewriting: the shared styling there was migrated
# out of three builders and verified by rebuilding each and comparing paragraph
# style, spacing, alignment, every run's text/bold/italic/size/colour, table
# dimensions, column widths and cell alignment. That is a lot of settled detail to
# re-derive for a fourth project.
#
# Copied 2026-08-14. Unmodified below this header.
# ============================================================================

"""Shared Word styling for this repo's generated reports.

`build_gauge_report.py` and `build_methods_doc.py` each grew their own copy of
these helpers and `build_report.py` would have been a third, so they live here
instead. All three builders now use this module.

The migration of the two older builders was required to leave their output
untouched, and was checked that way: each was run before and after, and the
documents compared on paragraph style, spacing, alignment, every run's
text/bold/italic/size/colour, table dimensions, column widths and cell
alignment. Both came out identical, which is why the committed `.docx` files
were not regenerated — there was nothing to regenerate.

Their three genuine behavioural differences survive as arguments to
`new_document` rather than as forks of the code; see its docstring.

Usage:

    S = new_document(FIGS_DIR)
    doc, para, table = S.doc, S.para, S.table
    S.h1('1.  Executive summary')

Everything writes into `S.doc`; call `S.doc.save(path)` when finished.
"""
import datetime as dt
import os
from types import SimpleNamespace

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# House palette. INK is body text, SOFT is captions and notes, ACCENT is
# headings and table chrome, WARN is the amber used for caveat callouts.
INK = RGBColor(0x16, 0x21, 0x2B)
SOFT = RGBColor(0x5B, 0x6B, 0x7A)
ACCENT = RGBColor(0x1F, 0x3B, 0x60)
WARN = RGBColor(0x9A, 0x64, 0x00)

HEADER_FILL = '1F3B60'
STRIPE_FILL = 'F2F5F8'
CALLOUT_FILL = 'EEF3F9'
CALLOUT_WARN_FILL = 'FDF3E0'

# (style name, pt, colour, space-before pt, space-after pt, page-break-before)
DEFAULT_HEADINGS = (
    ('Heading 1', 16.0, ACCENT, 18, 8, True),
    ('Heading 2', 12.5, ACCENT, 14, 5, False),
    ('Heading 3', 11.0, INK, 11, 4, False),
)

# matplotlib rcParams that make figures match the document. Import and pass to
# plt.rcParams.update() — kept here so figure styling cannot drift from text.
PLOT_RC = {
    'font.family': 'Arial', 'font.size': 9, 'axes.titlesize': 10,
    'axes.titleweight': 'bold', 'axes.edgecolor': '#5b6b7a',
    'axes.labelcolor': '#16212b', 'text.color': '#16212b',
    'xtick.color': '#5b6b7a', 'ytick.color': '#5b6b7a', 'axes.grid': True,
    'grid.color': '#dbe2e8', 'grid.linewidth': .7, 'legend.fontsize': 8,
    'legend.frameon': False, 'figure.dpi': 200, 'savefig.dpi': 200,
    'savefig.bbox': 'tight',
}
C_MEAS, C_MODEL, C_ALT = '#1f77b4', '#c0504d', '#ed7d31'
C_2022, C_WARN = '#4c9a2a', '#9a6400'


DATE_FMT = '%d %B %Y'


def build_date():
    """The date stamped on a generated document.

    Overridable, because otherwise "rebuild it and check nothing changed" stops
    being a usable signal the day after a commit: every Word document here
    carries a date on its title page, so a rebuild the next morning differs by
    one line for no real reason.

        SOURCE_DATE=2026-08-09 python tools/build_report.py
        SOURCE_DATE_EPOCH=1786060800 python tools/build_report.py

    `SOURCE_DATE` takes an ISO date and is the readable one. `SOURCE_DATE_EPOCH`
    is the reproducible-builds convention and takes a Unix timestamp, read in
    UTC. Neither set means today.

    A malformed value raises rather than falling back to today: a silent
    fallback on a typo'd variable would reintroduce exactly the spurious diff
    this exists to remove, and would do it invisibly.
    """
    iso = os.environ.get('SOURCE_DATE')
    if iso and iso.strip():
        try:
            return dt.date.fromisoformat(iso.strip())
        except ValueError:
            raise ValueError(f'SOURCE_DATE={iso!r} is not an ISO date '
                             f'(expected YYYY-MM-DD)') from None
    epoch = os.environ.get('SOURCE_DATE_EPOCH')
    if epoch and epoch.strip():
        try:
            return dt.datetime.fromtimestamp(int(epoch.strip()),
                                             dt.timezone.utc).date()
        except (ValueError, OverflowError, OSError):
            raise ValueError(f'SOURCE_DATE_EPOCH={epoch!r} is not a Unix '
                             f'timestamp') from None
    return dt.date.today()


def build_date_str(fmt=DATE_FMT):
    """`build_date()` in the house format — the string the title pages print."""
    return build_date().strftime(fmt)


def check_table_widths(doc, text_width_in=6.5, tol=0.02):
    """Fail the build if any table is wider than the text column.

    Column-wrap defects are invisible in the builder source and only show up
    when the document is rendered — this repo has paid for that lesson more
    than once. Cheap to check mechanically, so it is checked on every build.
    """
    bad = []
    for i, t in enumerate(doc.tables):
        w = sum((c.width.inches if c.width else 0.0) for c in t.rows[0].cells)
        if w > text_width_in + tol:
            head = ' | '.join(c.text for c in t.rows[0].cells)[:70]
            bad.append(f'  table {i}: {w:.2f}in > {text_width_in}in — {head}')
    if bad:
        raise ValueError(
            'these tables are wider than the text column and will wrap:\n'
            + '\n'.join(bad))


def new_document(figs_dir=None, headings=DEFAULT_HEADINGS, body_pt=10.0,
                 right_from=1, warn_prefix=False, callout_spacer=6):
    """A Document with this repo's page setup and styles, plus bound helpers.

    The three behaviour knobs exist because the documents genuinely differ, and
    migrating them onto this module had to leave their output byte-identical:

      right_from      default first right-aligned table column. The methods
                      report left-aligns everything (99); the others right-align
                      from column 1.
      warn_prefix     colour table cells starting '⚠' or '†' amber. The methods
                      report does; the gauge report and the efficiency report
                      do not.
      callout_spacer  points of trailing space after a callout, or None for no
                      spacer paragraph at all. The efficiency report wants one;
                      the two older documents do not.
    """
    doc = Document()
    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    st = doc.styles['Normal']
    st.font.name, st.font.size, st.font.color.rgb = 'Arial', Pt(body_pt), INK
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    for name, size, color, before, after, brk in headings:
        h = doc.styles[name]
        h.font.name, h.font.size, h.font.bold = 'Arial', Pt(size), True
        h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(before)
        h.paragraph_format.space_after = Pt(after)
        h.paragraph_format.keep_with_next = True
        h.paragraph_format.page_break_before = brk

    def para(t='', size=body_pt, bold=False, italic=False, color=None,
             align=None, after=6):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        if t:
            r = p.add_run(t)
            r.font.size, r.bold, r.italic = Pt(size), bold, italic
            r.font.name = 'Arial'
            if color is not None:
                r.font.color.rgb = color
        return p

    def rich(parts, size=body_pt, after=6, align=None):
        """A paragraph of (text, bold) pairs — for a bolded lead-in."""
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        for text, bold in parts:
            r = p.add_run(text)
            r.bold, r.font.size, r.font.name = bold, Pt(size), 'Arial'
        return p

    def mono(t, note=None):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2 if note else 8)
        r = p.add_run(t)
        r.font.name, r.font.size, r.bold = 'Consolas', Pt(9.5), True
        if note:
            q = doc.add_paragraph()
            q.paragraph_format.left_indent = Inches(0.3)
            q.paragraph_format.space_after = Pt(8)
            rr = q.add_run(note)
            rr.font.size, rr.italic, rr.font.name = Pt(8.5), True, 'Arial'
            rr.font.color.rgb = SOFT

    def bullets(items, size=body_pt):
        for it in items:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            if isinstance(it, tuple):
                a = p.add_run(it[0])
                a.bold, a.font.size, a.font.name = True, Pt(size), 'Arial'
                b = p.add_run(it[1])
                b.font.size, b.font.name = Pt(size), 'Arial'
            else:
                r = p.add_run(it)
                r.font.size, r.font.name = Pt(size), 'Arial'

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        el = OxmlElement('w:shd')
        el.set(qn('w:val'), 'clear')
        el.set(qn('w:color'), 'auto')
        el.set(qn('w:fill'), hexc)
        tcPr.append(el)

    def table(headers, rows, widths, note=None, size=8.5, right_from=right_from,
              flag_rows=()):
        """`flag_rows` are zero-based row indices to tint amber (caveats)."""
        t = doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        for i, w in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Inches(w)
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]
            c.width, c.text = Inches(widths[i]), ''
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if i >= right_from:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(h)
            r.bold, r.font.size, r.font.name = True, Pt(size), 'Arial'
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade(c, HEADER_FILL)
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].width, cells[i].text = Inches(widths[i]), ''
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(2)
                if i >= right_from:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run(str(v))
                r.font.size, r.font.name = Pt(size), 'Arial'
                if warn_prefix and str(v).startswith(('⚠', '†')):
                    r.font.color.rgb = WARN
            if ri in flag_rows:
                for c in cells:
                    shade(c, CALLOUT_WARN_FILL)
            elif ri % 2 == 1:
                for c in cells:
                    shade(c, STRIPE_FILL)
        if note:
            para(note, size=8, italic=True, color=SOFT, after=10)
        else:
            para('', after=6)
        return t

    def callout(title, text, color=ACCENT):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
        c = t.rows[0].cells[0]
        c.width, c.text = Inches(6.3), ''
        shade(c, CALLOUT_FILL if color == ACCENT else CALLOUT_WARN_FILL)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title)
        r.bold, r.font.size = True, Pt(9.5)
        r.font.name, r.font.color.rgb = 'Arial', color
        p2 = c.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r2 = p2.add_run(text)
        r2.font.size, r2.font.name = Pt(9), 'Arial'
        if callout_spacer is not None:
            para('', after=callout_spacer)

    def figure(name, caption, width=6.2):
        if figs_dir is None:
            raise RuntimeError('new_document() was given no figs_dir')
        path = figs_dir / name
        if not path.exists():
            raise FileNotFoundError(
                f'figure asset missing: {path}. Figures are either regenerated '
                f'by the builder or carried as versioned assets — a missing one '
                f'is a broken build, not a blank page.')
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(caption, size=8.5, italic=True, color=SOFT,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

    def heading(text, level):
        doc.add_heading(text, level=level)

    return SimpleNamespace(
        doc=doc, para=para, rich=rich, mono=mono, bullets=bullets, shade=shade,
        table=table, callout=callout, figure=figure,
        h1=lambda t: heading(t, 1), h2=lambda t: heading(t, 2),
        h3=lambda t: heading(t, 3),
        CENTER=WD_ALIGN_PARAGRAPH.CENTER, RIGHT=WD_ALIGN_PARAGRAPH.RIGHT,
        INK=INK, SOFT=SOFT, ACCENT=ACCENT, WARN=WARN)
